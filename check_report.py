"""
ECU报告DID件号自动检查脚本

用法: python check_report.py
- 从 part_number/part_number.xlsx 读取标准DID件号映射
- 自动扫描 template/ 下所有 Excel/Word 报告
- 通过件号前缀自动探测DID归属，比对实际值 vs 标准值
- 输出检查报告到 check_result/
"""

import os
import re
import glob
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document


# ─── 数据结构 ───

@dataclass
class CheckResult:
    file_name: str
    sheet_name: str
    location: str
    did: str
    expected: str
    actual: str
    status: str  # OK / NG / MISSING / UNKNOWN


# ─── 1. PartNumberLoader ───

CONTROLLER_NAMES = ["FLSPU", "RLSPU", "RRSPU", "FLSMU", "FRSMU"]

def load_part_numbers(xlsx_path: str) -> tuple[dict[str, dict[str, str]], dict[str, str]]:
    """
    从 part_number.xlsx 加载所有控制器的标准件号映射。
    返回:
        all_part_numbers: {"FLSPU": {"F1A0": "S000004325006", ...}, "RLSPU": {...}, ...}
        pn_prefix_to_did: {"S000004325": "F1A0", ...}  # 合并所有控制器的前缀映射
    """
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active

    # 从表头读取控制器名（H-L列，索引7-11），fallback到硬编码
    header_ctrls = {}
    for col_idx in range(8, 13):  # H=8, I=9, J=10, K=11, L=12
        val = ws.cell(row=1, column=col_idx).value
        if val:
            ctrl = str(val).strip()
            header_ctrls[col_idx] = ctrl

    # 如果表头读取失败，使用默认映射
    ctrl_cols = {}
    default_map = {8: "FLSPU", 9: "RLSPU", 10: "RRSPU", 11: "FLSMU", 12: "FRSMU"}
    for col_idx, default_name in default_map.items():
        ctrl_cols[col_idx] = header_ctrls.get(col_idx, default_name)

    # 读取每个控制器的件号
    all_part_numbers: dict[str, dict[str, str]] = {name: {} for name in ctrl_cols.values()}
    for row in ws.iter_rows(min_row=2, values_only=False):
        did_cell = row[2]  # C列 = DID
        if not did_cell.value:
            continue
        did = str(did_cell.value).strip()
        for col_idx, ctrl_name in ctrl_cols.items():
            pn_cell = row[col_idx - 1]  # row是0-indexed，col_idx是1-indexed
            if pn_cell.value and str(pn_cell.value).strip() not in ("", "-"):
                all_part_numbers[ctrl_name][did] = str(pn_cell.value).strip()

    wb.close()

    # 构建合并的件号前缀→DID反向映射（所有控制器共享，用于定位件号位置）
    pn_prefix_to_did = {}
    for ctrl_pns in all_part_numbers.values():
        for did, pn in ctrl_pns.items():
            if did == "F17F":
                pn_prefix_to_did[pn[:9].upper()] = did
            else:
                pn_prefix_to_did[pn[:10]] = did

    return all_part_numbers, pn_prefix_to_did


def resolve_controller(rel_path: str) -> list[str]:
    """
    从文件相对路径中识别控制器名称。
    返回控制器列表（多数情况为1个；FLSMU&FRSMU下无控制器名的文件返回2个）。
    """
    path_upper = rel_path.upper()

    # 特殊处理: FLSMU&FRSMU 文件夹 — 文件名可能单独含 FLSMU 或 FRSMU
    if "FLSMU" in path_upper and "FRSMU" in path_upper:
        # 路径同时包含两个关键词，检查文件名是否单独指定
        fname_upper = os.path.basename(path_upper)
        has_flsmu = "FLSMU" in fname_upper
        has_frsmu = "FRSMU" in fname_upper
        if has_flsmu and has_frsmu:
            return ["FLSMU", "FRSMU"]
        elif has_flsmu:
            return ["FLSMU"]
        elif has_frsmu:
            return ["FRSMU"]
        else:
            # 文件名不含控制器名 → 两个都关联
            return ["FLSMU", "FRSMU"]

    # 常规情况: 路径中包含单个控制器名
    matched = []
    for ctrl in CONTROLLER_NAMES:
        if ctrl in path_upper:
            matched.append(ctrl)
    return matched


# ─── 2. ExcelChecker ───

class ExcelChecker:
    def __init__(self, pn_prefix_to_did: dict[str, str]):
        self.pn_prefix_to_did = pn_prefix_to_did
        # DID代码正则
        self.did_pattern = re.compile(r'\b(F1[0-9A-Fa-f]{2}|F18E|F17F)\b')
        # UDS命令格式 "22 F1 xx"
        self.uds_pattern = re.compile(r'22\s+F1\s+([0-9A-Fa-f]{2})')
        # 关键词→DID映射（用于同行/同列左侧扫描）
        # 优先级从上到下，长关键词优先匹配
        self.keyword_to_did = {
            # 英文标签（来自part_number.xlsx和各模板）
            "SWDI-SAF1": "F103",    # SWDI-SFA1 → APP诊断数据库 → F103
            "SWDI-SFA1": "F103",
            "SWDI-PBL": "F104",     # PBL诊断数据库零件号 → F104
            "SWDI-SBL": "F105",     # SBL诊断数据库零件号 → F105
            "SFA1": "F1A0",         # 应用层软件 → F1A0
            "SAF1": "F1A0",         # 同SFA1
            "SFBL": "F102",         # Secondary boot loader → F102
            "SBL": "F102",          # Secondary boot loader → F102
            "PBL": "F180",          # Primary boot loader → F180
            "SPBL": "F180",         # 同PBL
            "HWBN": "F193",         # 控制器硬件 → F193
            "DU PN": "F18E",        # 控制器 → F18E
            "Application": "F1A0",
            # 中文标签
            "ECU总成号": "F18E",
            "ECU总成": "F18E",
            "硬件号": "F193",
            "硬件版本": "F193",
            "初始引导程序号": "F180",
            "初始引导程序": "F180",
            "应用层软件号": "F1A0",
            "应用层软件": "F1A0",
            "PBL诊断数据库": "F104",
            "SBL诊断数据库": "F105",
            "APP诊断数据库": "F103",
        }

    def check_file(self, filepath: str, rel_path: str = "", part_numbers: dict[str, str] = None) -> list[CheckResult]:
        results = []
        fname = rel_path or os.path.basename(filepath)
        try:
            wb = openpyxl.load_workbook(filepath, data_only=True)
        except Exception as e:
            print(f"  [WARN] 无法打开 {fname}: {e}")
            return results

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            results.extend(self.check_sheet(ws, fname, sheet_name, part_numbers))

        wb.close()
        return results

    def check_sheet(self, ws, file_name: str, sheet_name: str, part_numbers: dict[str, str] = None) -> list[CheckResult]:
        if part_numbers is None:
            part_numbers = {}
        results = []
        # 记录每个DID已找到的件号（每sheet每个DID只报一次）
        found_dids = {}

        max_row = ws.max_row or 0
        max_col = ws.max_column or 0

        # 判断是否是有多行历史版本的sheet（如PN MAP、HW/SW changelog）
        # 如果是，只检查最后一行件号数据
        latest_row_only = self._detect_latest_row_only(ws, sheet_name, max_row, max_col)

        for row in range(1, max_row + 1):
            # 如果是仅检查最新行模式，跳过非最新行
            if latest_row_only is not None and row != latest_row_only:
                continue

            for col in range(1, max_col + 1):
                cell = ws.cell(row=row, column=col)
                val = cell.value
                if val is None:
                    continue
                val_str = str(val).strip()

                # 检查是否以任一件号前缀开头
                matched_prefix = self._match_pn_prefix(val_str)
                if not matched_prefix:
                    continue

                # 尝试判断DID归属
                did = self._detect_did_from_context(ws, row, col, matched_prefix)
                cell_addr = cell.coordinate

                if did is None:
                    # 件号前缀本身可以反查DID
                    did = self.pn_prefix_to_did.get(matched_prefix, None)

                if did is None:
                    results.append(CheckResult(
                        file_name=file_name, sheet_name=sheet_name,
                        location=cell_addr, did="UNKNOWN",
                        expected="-", actual=val_str, status="UNKNOWN"
                    ))
                    continue

                expected = part_numbers.get(did, "")
                actual = val_str
                # 对于F17F，实际值可能是hex带空格
                if did == "F17F":
                    actual_clean = actual.replace(" ", "").upper()
                    expected_clean = expected.upper()
                    status = "OK" if actual_clean == expected_clean else "NG"
                else:
                    status = "OK" if actual == expected else "NG"

                # 每个DID在每个sheet只记录一次（优先非空值）
                if did not in found_dids or found_dids[did].actual == "-":
                    found_dids[did] = CheckResult(
                        file_name=file_name, sheet_name=sheet_name,
                        location=cell_addr, did=did,
                        expected=expected, actual=actual, status=status
                    )

        results.extend(found_dids.values())
        return results

    def _match_pn_prefix(self, val: str) -> Optional[str]:
        """检查值是否以任一件号前缀开头"""
        val_upper = val.upper()
        for prefix in self.pn_prefix_to_did:
            if val_upper.startswith(prefix):
                return prefix
        return None

    def _detect_did_from_context(self, ws, row: int, col: int, pn_prefix: str) -> Optional[str]:
        """通过上下文判断件号的DID归属"""
        # 策略1: 同列上方搜索UDS命令 "22 F1 xx"
        for r in range(max(1, row - 5), row):
            cell_val = ws.cell(row=r, column=col).value
            if cell_val:
                m = self.uds_pattern.search(str(cell_val))
                if m:
                    return f"F1{m.group(1).upper()}"

        # 策略2: 同行左侧搜索DID代码（最多左扫8列）
        for c in range(max(1, col - 8), col):
            cell_val = ws.cell(row=row, column=c).value
            if cell_val:
                val_str = str(cell_val).strip()
                m = self.did_pattern.search(val_str)
                if m:
                    return m.group(1).upper()

        # 策略3: 同行左侧搜索关键词（按关键词长度降序，长关键词优先）
        sorted_keywords = sorted(self.keyword_to_did.items(), key=lambda x: len(x[0]), reverse=True)
        for c in range(max(1, col - 8), col):
            cell_val = ws.cell(row=row, column=c).value
            if cell_val:
                val_str = str(cell_val).strip().upper()
                for keyword, did in sorted_keywords:
                    if keyword.upper() in val_str:
                        return did

        # 策略4: 同列上方搜索关键词（最多上扫3行）
        for r in range(max(1, row - 3), row):
            cell_val = ws.cell(row=r, column=col).value
            if cell_val:
                val_str = str(cell_val).strip().upper()
                for keyword, did in sorted_keywords:
                    if keyword.upper() in val_str:
                        return did

        return None

    def _detect_latest_row_only(self, ws, sheet_name: str, max_row: int, max_col: int) -> Optional[int]:
        """
        检测是否是有多行历史版本的sheet。
        如果是，返回应该检查的最后一行号；否则返回None表示检查所有行。
        """
        sheet_upper = sheet_name.upper()

        # PN MAP sheet: 找最后一行有E4U版本标记的行
        if "PN MAP" in sheet_upper or sheet_upper == "PN MAP":
            last_data_row = None
            for row in range(1, max_row + 1):
                a_val = ws.cell(row=row, column=1).value
                if a_val and re.search(r'E4U\d', str(a_val)):
                    last_data_row = row
            return last_data_row

        # HW changelog: 找最后一行有日期数据的行
        if "HW" in sheet_upper and "CHANGELOG" in sheet_upper:
            last_data_row = None
            for row in range(1, max_row + 1):
                a_val = ws.cell(row=row, column=1).value
                if a_val and re.match(r'\d{4}\.\d{2}\.\d{2}', str(a_val)):
                    last_data_row = row
            return last_data_row

        # SW changelog: 找最后一行有E4U版本标记的行
        if "SW" in sheet_upper and "CHANGELOG" in sheet_upper:
            last_data_row = None
            for row in range(1, max_row + 1):
                a_val = ws.cell(row=row, column=1).value
                if a_val and re.search(r'E4U\d', str(a_val)):
                    last_data_row = row
            return last_data_row

        return None


# ─── 3. WordChecker ───

class WordChecker:
    def __init__(self, pn_prefix_to_did: dict[str, str]):
        self.pn_prefix_to_did = pn_prefix_to_did
        self.did_title_pattern = re.compile(r'^(F1[0-9A-Fa-f]{2}|F18E|F17F)\s*-+')
        self.version_pattern = re.compile(r'Version\s*=\s*(\S+)')
        self.hex_data_pattern = re.compile(r'原始数据\s*=\s*([\dA-Fa-f\s]+)')
        self.did_pattern = re.compile(r'\b(F1[0-9A-Fa-f]{2}|F18E|F17F)\b')

    def check_file(self, filepath: str, rel_path: str = "", part_numbers: dict[str, str] = None) -> list[CheckResult]:
        if part_numbers is None:
            part_numbers = {}
        fname = rel_path or os.path.basename(filepath)
        try:
            doc = Document(filepath)
        except Exception as e:
            print(f"  [WARN] 无法打开 {fname}: {e}")
            return []

        results = []
        # Part A: 表格检查
        results.extend(self._check_tables(doc, fname, part_numbers))
        # Part B: 段落检查
        results.extend(self._check_paragraphs(doc, fname, part_numbers))
        return results

    def _check_tables(self, doc, file_name: str, part_numbers: dict[str, str]) -> list[CheckResult]:
        results = []

        for table_idx, table in enumerate(doc.tables):
            # 策略1: 找含DID列的表
            did_col = None
            pn_col = None
            header_row_idx = None

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip().upper()
                    if "DID" in text and "对应" in text:
                        did_col = col_idx
                        header_row_idx = row_idx
                    elif text == "DID":
                        did_col = col_idx
                        header_row_idx = row_idx
                    # 件号列通常在最左或标记为"件号"
                    if "件号" in text or "PART" in text.upper():
                        pn_col = col_idx

                if did_col is not None:
                    break

            if did_col is not None:
                # 遍历DID列的每一行
                for row_idx in range((header_row_idx or 0) + 1, len(table.rows)):
                    row = table.rows[row_idx]
                    did_text = row.cells[did_col].text.strip()

                    # 提取DID代码
                    m = self.did_pattern.search(did_text)
                    if not m:
                        continue
                    did = m.group(1).upper()

                    if did not in part_numbers:
                        continue

                    # 获取件号值
                    if pn_col is not None:
                        actual = row.cells[pn_col].text.strip()
                    else:
                        # 件号通常在DID列的前一列或后一列
                        actual = ""
                        for ci in range(len(row.cells)):
                            if ci != did_col:
                                cell_text = row.cells[ci].text.strip()
                                # 检查是否像件号
                                for prefix in self.pn_prefix_to_did:
                                    if cell_text.upper().startswith(prefix):
                                        actual = cell_text
                                        break
                                if actual:
                                    break

                    if not actual:
                        continue

                    expected = part_numbers[did]
                    status = "OK" if actual == expected else "NG"
                    results.append(CheckResult(
                        file_name=file_name, sheet_name=f"表格{table_idx}",
                        location=f"行{row_idx}", did=did,
                        expected=expected, actual=actual, status=status
                    ))

            # 策略2: 件号前缀扫描表格单元格
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    matched_prefix = self._match_pn_prefix(text)
                    if not matched_prefix:
                        continue

                    did = self.pn_prefix_to_did.get(matched_prefix)
                    if not did or did not in part_numbers:
                        continue

                    # 检查这个DID是否已经被策略1记录
                    already_found = any(
                        r.did == did and r.file_name == file_name
                        for r in results
                    )
                    if already_found:
                        continue

                    expected = part_numbers[did]
                    status = "OK" if text == expected else "NG"
                    results.append(CheckResult(
                        file_name=file_name, sheet_name=f"表格{table_idx}",
                        location=f"R{row_idx}C{col_idx}", did=did,
                        expected=expected, actual=text, status=status
                    ))

        return results

    def _check_paragraphs(self, doc, file_name: str, part_numbers: dict[str, str]) -> list[CheckResult]:
        results = []
        paragraphs = doc.paragraphs

        # 扫描段落找DID标题和对应的Version
        i = 0
        while i < len(paragraphs):
            text = paragraphs[i].text.strip()
            m = self.did_title_pattern.match(text)
            if m:
                did = m.group(1).upper()
                # 在后续段落中搜索Version或hex数据
                version_found = None
                hex_data_found = None
                location_start = i

                for j in range(i + 1, min(i + 10, len(paragraphs))):
                    next_text = paragraphs[j].text.strip()

                    # 如果遇到下一个DID标题，停止搜索
                    if j > i + 1 and self.did_title_pattern.match(next_text):
                        break

                    vm = self.version_pattern.search(next_text)
                    if vm:
                        version_found = vm.group(1)

                    hm = self.hex_data_pattern.search(next_text)
                    if hm:
                        hex_data_found = hm.group(1).replace(" ", "").upper()

                # 生成CheckResult
                if did in part_numbers:
                    expected = part_numbers[did]

                    if did == "F17F" and hex_data_found:
                        actual = hex_data_found
                        expected_clean = expected.upper()
                        status = "OK" if actual == expected_clean else "NG"
                        results.append(CheckResult(
                            file_name=file_name, sheet_name="实测过程",
                            location=f"段落{location_start + 1}-{j}", did=did,
                            expected=expected, actual=actual, status=status
                        ))
                    elif version_found:
                        actual = version_found
                        status = "OK" if actual == expected else "NG"
                        results.append(CheckResult(
                            file_name=file_name, sheet_name="实测过程",
                            location=f"段落{location_start + 1}-{j}", did=did,
                            expected=expected, actual=actual, status=status
                        ))
            i += 1

        # 检查MISSING的DID（在实测过程中出现但没找到值的，或完全没出现的）
        found_in_paragraphs = {r.did for r in results}

        # 只有当实测过程区域确实存在时才检查MISSING
        # 判断依据：之前循环是否找到过任何DID标题
        has_test_section = len(found_in_paragraphs) > 0

        if has_test_section:
            for did, expected in part_numbers.items():
                if did not in found_in_paragraphs:
                    results.append(CheckResult(
                        file_name=file_name, sheet_name="实测过程",
                        location="-", did=did,
                        expected=expected, actual="-", status="MISSING"
                    ))

        return results

    def _match_pn_prefix(self, val: str) -> Optional[str]:
        val_upper = val.upper()
        for prefix in self.pn_prefix_to_did:
            if val_upper.startswith(prefix):
                return prefix
        return None


# ─── 4. ReportGenerator ───

class ReportGenerator:
    HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
    NG_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    MISSING_FILL = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    UNKNOWN_FILL = PatternFill(start_color="B4C6E7", end_color="B4C6E7", fill_type="solid")
    OK_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    THIN_BORDER = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    DETAIL_HEADERS = ["序号", "文件名", "Sheet/区域", "位置", "DID", "期望件号", "实际件号", "结果"]
    COL_WIDTHS = [6, 40, 18, 12, 8, 18, 18, 10]

    def generate(self, results_by_controller: dict[str, list[CheckResult]], output_path: str):
        wb = openpyxl.Workbook()
        first_sheet = True

        for ctrl, results in results_by_controller.items():
            if not results:
                continue
            if first_sheet:
                ws = wb.active
                ws.title = ctrl
                first_sheet = False
            else:
                ws = wb.create_sheet(ctrl)
            self._write_detail_sheet(ws, results)

        # 汇总统计 sheet
        ws_summary = wb.create_sheet("汇总统计") if not first_sheet else wb.active
        if first_sheet:
            ws_summary.title = "汇总统计"
        self._write_summary_sheet(ws_summary, results_by_controller)

        wb.save(output_path)
        print(f"\n检查报告已生成: {output_path}")

    def _write_detail_sheet(self, ws, results: list[CheckResult]):
        # 表头
        for col, header in enumerate(self.DETAIL_HEADERS, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = self.HEADER_FILL
            cell.font = self.HEADER_FONT
            cell.alignment = Alignment(horizontal='center')
            cell.border = self.THIN_BORDER

        for idx, r in enumerate(results, 1):
            row_data = [idx, r.file_name, r.sheet_name, r.location, r.did, r.expected, r.actual, r.status]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(row=idx + 1, column=col, value=val)
                cell.border = self.THIN_BORDER
                cell.alignment = Alignment(horizontal='center' if col in (1, 5, 8) else 'left')

            # 根据状态标色
            if r.status == "NG":
                for c in range(1, 9):
                    ws.cell(row=idx + 1, column=c).fill = self.NG_FILL
            elif r.status == "MISSING":
                for c in range(1, 9):
                    ws.cell(row=idx + 1, column=c).fill = self.MISSING_FILL
            elif r.status == "UNKNOWN":
                for c in range(1, 9):
                    ws.cell(row=idx + 1, column=c).fill = self.UNKNOWN_FILL
            elif r.status == "OK":
                ws.cell(row=idx + 1, column=8).fill = self.OK_FILL

        for i, w in enumerate(self.COL_WIDTHS, 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    def _write_summary_sheet(self, ws, results_by_controller: dict[str, list[CheckResult]]):
        row = 1
        # 总表头
        summary_headers = ["控制器", "总检查数", "OK", "NG", "MISSING", "UNKNOWN"]
        for col, h in enumerate(summary_headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.fill = self.HEADER_FILL
            cell.font = self.HEADER_FONT
            cell.alignment = Alignment(horizontal='center')
            cell.border = self.THIN_BORDER

        row = 2
        for ctrl, results in results_by_controller.items():
            if not results:
                continue
            ok_count = sum(1 for r in results if r.status == "OK")
            ng_count = sum(1 for r in results if r.status == "NG")
            miss_count = sum(1 for r in results if r.status == "MISSING")
            unk_count = sum(1 for r in results if r.status == "UNKNOWN")
            for col, val in enumerate([ctrl, len(results), ok_count, ng_count, miss_count, unk_count], 1):
                cell = ws.cell(row=row, column=col, value=val)
                cell.border = self.THIN_BORDER
                cell.alignment = Alignment(horizontal='center')
            row += 1

        for i, w in enumerate([14, 12, 8, 8, 10, 10], 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

        # NG详情（所有控制器合并）
        row += 1
        all_ng = [(ctrl, r) for ctrl, rs in results_by_controller.items() for r in rs if r.status == "NG"]
        if all_ng:
            ws.cell(row=row, column=1, value="NG 详情").font = Font(bold=True, size=12)
            row += 1
            ng_headers = ["控制器", "文件名", "Sheet/区域", "DID", "期望件号", "实际件号"]
            for col, h in enumerate(ng_headers, 1):
                cell = ws.cell(row=row, column=col, value=h)
                cell.fill = self.NG_FILL
                cell.font = Font(bold=True)
                cell.border = self.THIN_BORDER
            row += 1
            for ctrl, r in all_ng:
                for col, val in enumerate([ctrl, r.file_name, r.sheet_name, r.did, r.expected, r.actual], 1):
                    cell = ws.cell(row=row, column=col, value=val)
                    cell.border = self.THIN_BORDER
                row += 1

        # MISSING详情
        row += 1
        all_miss = [(ctrl, r) for ctrl, rs in results_by_controller.items() for r in rs if r.status == "MISSING"]
        if all_miss:
            ws.cell(row=row, column=1, value="MISSING 详情").font = Font(bold=True, size=12)
            row += 1
            miss_headers = ["控制器", "文件名", "Sheet/区域", "DID", "期望件号"]
            for col, h in enumerate(miss_headers, 1):
                cell = ws.cell(row=row, column=col, value=h)
                cell.fill = self.MISSING_FILL
                cell.font = Font(bold=True)
                cell.border = self.THIN_BORDER
            row += 1
            for ctrl, r in all_miss:
                for col, val in enumerate([ctrl, r.file_name, r.sheet_name, r.did, r.expected], 1):
                    cell = ws.cell(row=row, column=col, value=val)
                    cell.border = self.THIN_BORDER
                row += 1


# ─── 5. Main ───

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    pn_path = os.path.join(base_dir, "part_number", "part_number.xlsx")
    template_dir = os.path.join(base_dir, "template")
    result_dir = os.path.join(base_dir, "check_result")

    # 检查路径
    if not os.path.exists(pn_path):
        print(f"错误: 找不到标准件号文件 {pn_path}")
        return
    if not os.path.exists(template_dir):
        print(f"错误: 找不到模板目录 {template_dir}")
        return

    # 1. 加载标准件号
    print("=" * 60)
    print("ECU报告DID件号自动检查")
    print("=" * 60)
    print("\n[1/4] 加载标准件号...")
    all_part_numbers, pn_prefix_to_did = load_part_numbers(pn_path)
    for ctrl, pns in all_part_numbers.items():
        print(f"  {ctrl}: {len(pns)} 个DID")

    # 2. 扫描文件（递归遍历所有子文件夹，跳过Office临时文件）
    print("\n[2/4] 扫描 template/ 目录...")
    xlsx_files = glob.glob(os.path.join(template_dir, "**", "*.xlsx"), recursive=True)
    docx_files = glob.glob(os.path.join(template_dir, "**", "*.docx"), recursive=True)
    all_files = [f for f in (xlsx_files + docx_files) if not os.path.basename(f).startswith("~$")]
    print(f"  找到 {len(all_files)} 个文件")

    # 3. 检查每个文件（按控制器分组）
    print("\n[3/4] 开始检查...")
    results_by_controller: dict[str, list[CheckResult]] = {ctrl: [] for ctrl in all_part_numbers}
    excel_checker = ExcelChecker(pn_prefix_to_did)
    word_checker = WordChecker(pn_prefix_to_did)

    for filepath in all_files:
        rel_path = os.path.relpath(filepath, template_dir)
        controllers = resolve_controller(rel_path)
        if not controllers:
            print(f"  [SKIP] 无法识别控制器: {rel_path}")
            continue

        ctrl_label = "/".join(controllers)
        print(f"  检查: {rel_path}  [{ctrl_label}]")

        for ctrl in controllers:
            pn = all_part_numbers.get(ctrl, {})
            if filepath.endswith(".xlsx"):
                results = excel_checker.check_file(filepath, rel_path, pn)
            elif filepath.endswith(".docx"):
                results = word_checker.check_file(filepath, rel_path, pn)
            else:
                continue

            if results:
                results_by_controller[ctrl].extend(results)
                ok = sum(1 for r in results if r.status == "OK")
                ng = sum(1 for r in results if r.status == "NG")
                miss = sum(1 for r in results if r.status == "MISSING")
                unk = sum(1 for r in results if r.status == "UNKNOWN")
                print(f"    [{ctrl}] → OK:{ok} NG:{ng} MISSING:{miss} UNKNOWN:{unk}")
            else:
                print(f"    [{ctrl}] → 未发现件号（跳过）")

    # 4. 生成报告
    print("\n[4/4] 生成检查报告...")
    os.makedirs(result_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(result_dir, f"检查报告_{timestamp}.xlsx")
    report_gen = ReportGenerator()
    report_gen.generate(results_by_controller, output_path)

    # 终端摘要
    print("\n" + "=" * 60)
    print("检查结果摘要")
    print("=" * 60)

    for ctrl, results in results_by_controller.items():
        if not results:
            continue
        ok_list = [r for r in results if r.status == "OK"]
        ng_list = [r for r in results if r.status == "NG"]
        miss_list = [r for r in results if r.status == "MISSING"]
        unk_list = [r for r in results if r.status == "UNKNOWN"]

        print(f"\n── {ctrl} ({len(results)} 项) ──")
        print(f"  OK:      {len(ok_list)} 项")
        print(f"  NG:      {len(ng_list)} 项")
        print(f"  MISSING: {len(miss_list)} 项")
        print(f"  UNKNOWN: {len(unk_list)} 项")

        if ng_list:
            print(f"  [NG] 错误详情:")
            for r in ng_list:
                print(f"    [{r.file_name}] {r.sheet_name} {r.location} | {r.did}: {r.actual} (期望: {r.expected})")

        if miss_list:
            print(f"  [MISSING] 缺失详情:")
            for r in miss_list:
                print(f"    [{r.file_name}] {r.sheet_name} {r.location} | {r.did}: 期望 {r.expected}")

        if unk_list:
            print(f"  [UNKNOWN] 未知详情:")
            for r in unk_list:
                print(f"    [{r.file_name}] {r.sheet_name} {r.location} | 值: {r.actual}")

    total_results = [r for rs in results_by_controller.values() for r in rs]
    has_issues = any(r.status in ("NG", "MISSING") for r in total_results)
    status = "PASS" if not has_issues else "ISSUES FOUND"
    print(f"\n[{status}] 检查完成!")


if __name__ == "__main__":
    main()
