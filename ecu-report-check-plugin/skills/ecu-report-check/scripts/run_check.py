"""
ECU报告DID件号自动检查脚本 (OpenClaw Skill 版本)

用法:
  CLI:  python run_check.py --files a.xlsx b.docx --controllers FLSMU --pn-file part_number.xlsx
  API:  from run_check import run_check
        result = run_check(files=[...], controllers=[...], pn_file=...)

输出 JSON 到 stdout。
"""

import os
import re
import sys
import json
import argparse
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document


# ─── 数据结构 ───

@dataclass
class CheckResult:
    file_name: str
    sheet_name: str
    location: str
    did: str
    expected: str
    actual: str
    status: str  # OK / NG / MISSING / UNKNOWN

    def to_dict(self) -> dict:
        return {
            "file_name": self.file_name,
            "sheet_name": self.sheet_name,
            "location": self.location,
            "did": self.did,
            "expected": self.expected,
            "actual": self.actual,
            "status": self.status,
        }


CONTROLLER_NAMES = ["FLSPU", "RLSPU", "RRSPU", "FLSMU", "FRSMU"]


# ─── 1. PartNumberLoader ───

def load_part_numbers(xlsx_path: str) -> tuple[dict[str, dict[str, str]], dict[str, str]]:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active

    header_ctrls = {}
    for col_idx in range(8, 13):
        val = ws.cell(row=1, column=col_idx).value
        if val:
            ctrl = str(val).strip()
            header_ctrls[col_idx] = ctrl

    ctrl_cols = {}
    default_map = {8: "FLSPU", 9: "RLSPU", 10: "RRSPU", 11: "FLSMU", 12: "FRSMU"}
    for col_idx, default_name in default_map.items():
        ctrl_cols[col_idx] = header_ctrls.get(col_idx, default_name)

    all_part_numbers: dict[str, dict[str, str]] = {name: {} for name in ctrl_cols.values()}
    for row in ws.iter_rows(min_row=2, values_only=False):
        did_cell = row[2]
        if not did_cell.value:
            continue
        did = str(did_cell.value).strip()
        for col_idx, ctrl_name in ctrl_cols.items():
            pn_cell = row[col_idx - 1]
            if pn_cell.value and str(pn_cell.value).strip() not in ("", "-"):
                all_part_numbers[ctrl_name][did] = str(pn_cell.value).strip()

    wb.close()

    pn_prefix_to_did = {}
    for ctrl_pns in all_part_numbers.values():
        for did, pn in ctrl_pns.items():
            if did == "F17F":
                pn_prefix_to_did[pn[:9].upper()] = did
            else:
                pn_prefix_to_did[pn[:10]] = did

    return all_part_numbers, pn_prefix_to_did


def resolve_controller(rel_path: str) -> list[str]:
    path_upper = rel_path.upper()

    if "FLSMU" in path_upper and "FRSMU" in path_upper:
        fname_upper = os.path.basename(path_upper)
        has_flsmu = "FLSMU" in fname_upper
        has_frsmu = "FRSMU" in fname_upper
        if has_flsmu and has_frsmu:
            return ["FLSMU", "FRSMU"]
        elif has_flsmu:
            return ["FLSMU"]
        elif has_frsmu:
            return ["FRSMU"]
        else:
            return ["FLSMU", "FRSMU"]

    matched = []
    for ctrl in CONTROLLER_NAMES:
        if ctrl in path_upper:
            matched.append(ctrl)
    return matched


# ─── 2. ExcelChecker ───

class ExcelChecker:
    def __init__(self, pn_prefix_to_did: dict[str, str]):
        self.pn_prefix_to_did = pn_prefix_to_did
        self.did_pattern = re.compile(r'\b(F1[0-9A-Fa-f]{2}|F18E|F17F)\b')
        self.uds_pattern = re.compile(r'22\s+F1\s+([0-9A-Fa-f]{2})')
        self.keyword_to_did = {
            "SWDI-SAF1": "F103",
            "SWDI-SFA1": "F103",
            "SWDI-PBL": "F104",
            "SWDI-SBL": "F105",
            "SFA1": "F1A0",
            "SAF1": "F1A0",
            "SFBL": "F102",
            "SBL": "F102",
            "PBL": "F180",
            "SPBL": "F180",
            "HWBN": "F193",
            "DU PN": "F18E",
            "Application": "F1A0",
            "ECU总成号": "F18E",
            "ECU总成": "F18E",
            "硬件号": "F193",
            "硬件版本": "F193",
            "初始引导程序号": "F180",
            "初始引导程序": "F180",
            "应用层软件号": "F1A0",
            "应用层软件": "F1A0",
            "PBL诊断数据库": "F104",
            "SBL诊断数据库": "F105",
            "APP诊断数据库": "F103",
        }

    def check_file(self, filepath: str, rel_path: str = "", part_numbers: dict[str, str] = None) -> list[CheckResult]:
        results = []
        fname = rel_path or os.path.basename(filepath)
        try:
            wb = openpyxl.load_workbook(filepath, data_only=True)
        except Exception as e:
            return results

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            results.extend(self.check_sheet(ws, fname, sheet_name, part_numbers))

        wb.close()
        return results

    def check_sheet(self, ws, file_name: str, sheet_name: str, part_numbers: dict[str, str] = None) -> list[CheckResult]:
        if part_numbers is None:
            part_numbers = {}
        results = []
        found_dids = {}

        max_row = ws.max_row or 0
        max_col = ws.max_column or 0

        latest_row_only = self._detect_latest_row_only(ws, sheet_name, max_row, max_col)

        for row in range(1, max_row + 1):
            if latest_row_only is not None and row != latest_row_only:
                continue

            for col in range(1, max_col + 1):
                cell = ws.cell(row=row, column=col)
                val = cell.value
                if val is None:
                    continue
                val_str = str(val).strip()

                matched_prefix = self._match_pn_prefix(val_str)
                if not matched_prefix:
                    continue

                did = self._detect_did_from_context(ws, row, col, matched_prefix)
                cell_addr = cell.coordinate

                if did is None:
                    did = self.pn_prefix_to_did.get(matched_prefix, None)

                if did is None:
                    results.append(CheckResult(
                        file_name=file_name, sheet_name=sheet_name,
                        location=cell_addr, did="UNKNOWN",
                        expected="-", actual=val_str, status="UNKNOWN"
                    ))
                    continue

                expected = part_numbers.get(did, "")
                actual = val_str
                if did == "F17F":
                    actual_clean = actual.replace(" ", "").upper()
                    expected_clean = expected.upper()
                    status = "OK" if actual_clean == expected_clean else "NG"
                else:
                    status = "OK" if actual == expected else "NG"

                if did not in found_dids or found_dids[did].actual == "-":
                    found_dids[did] = CheckResult(
                        file_name=file_name, sheet_name=sheet_name,
                        location=cell_addr, did=did,
                        expected=expected, actual=actual, status=status
                    )

        results.extend(found_dids.values())
        return results

    def _match_pn_prefix(self, val: str) -> Optional[str]:
        val_upper = val.upper()
        for prefix in self.pn_prefix_to_did:
            if val_upper.startswith(prefix):
                return prefix
        return None

    def _detect_did_from_context(self, ws, row: int, col: int, pn_prefix: str) -> Optional[str]:
        for r in range(max(1, row - 5), row):
            cell_val = ws.cell(row=r, column=col).value
            if cell_val:
                m = self.uds_pattern.search(str(cell_val))
                if m:
                    return f"F1{m.group(1).upper()}"

        for c in range(max(1, col - 8), col):
            cell_val = ws.cell(row=row, column=c).value
            if cell_val:
                val_str = str(cell_val).strip()
                m = self.did_pattern.search(val_str)
                if m:
                    return m.group(1).upper()

        sorted_keywords = sorted(self.keyword_to_did.items(), key=lambda x: len(x[0]), reverse=True)
        for c in range(max(1, col - 8), col):
            cell_val = ws.cell(row=row, column=c).value
            if cell_val:
                val_str = str(cell_val).strip().upper()
                for keyword, did in sorted_keywords:
                    if keyword.upper() in val_str:
                        return did

        for r in range(max(1, row - 3), row):
            cell_val = ws.cell(row=r, column=col).value
            if cell_val:
                val_str = str(cell_val).strip().upper()
                for keyword, did in sorted_keywords:
                    if keyword.upper() in val_str:
                        return did

        return None

    def _detect_latest_row_only(self, ws, sheet_name: str, max_row: int, max_col: int) -> Optional[int]:
        sheet_upper = sheet_name.upper()

        if "PN MAP" in sheet_upper or sheet_upper == "PN MAP":
            last_data_row = None
            for row in range(1, max_row + 1):
                a_val = ws.cell(row=row, column=1).value
                if a_val and re.search(r'E4U\d', str(a_val)):
                    last_data_row = row
            return last_data_row

        if "HW" in sheet_upper and "CHANGELOG" in sheet_upper:
            last_data_row = None
            for row in range(1, max_row + 1):
                a_val = ws.cell(row=row, column=1).value
                if a_val and re.match(r'\d{4}\.\d{2}\.\d{2}', str(a_val)):
                    last_data_row = row
            return last_data_row

        if "SW" in sheet_upper and "CHANGELOG" in sheet_upper:
            last_data_row = None
            for row in range(1, max_row + 1):
                a_val = ws.cell(row=row, column=1).value
                if a_val and re.search(r'E4U\d', str(a_val)):
                    last_data_row = row
            return last_data_row

        return None


# ─── 3. WordChecker ───

class WordChecker:
    def __init__(self, pn_prefix_to_did: dict[str, str]):
        self.pn_prefix_to_did = pn_prefix_to_did
        self.did_title_pattern = re.compile(r'^(F1[0-9A-Fa-f]{2}|F18E|F17F)\s*-+')
        self.version_pattern = re.compile(r'Version\s*=\s*(\S+)')
        self.hex_data_pattern = re.compile(r'原始数据\s*=\s*([\dA-Fa-f\s]+)')
        self.did_pattern = re.compile(r'\b(F1[0-9A-Fa-f]{2}|F18E|F17F)\b')

    def check_file(self, filepath: str, rel_path: str = "", part_numbers: dict[str, str] = None) -> list[CheckResult]:
        if part_numbers is None:
            part_numbers = {}
        fname = rel_path or os.path.basename(filepath)
        try:
            doc = Document(filepath)
        except Exception as e:
            return []

        results = []
        results.extend(self._check_tables(doc, fname, part_numbers))
        results.extend(self._check_paragraphs(doc, fname, part_numbers))
        return results

    def _check_tables(self, doc, file_name: str, part_numbers: dict[str, str]) -> list[CheckResult]:
        results = []

        for table_idx, table in enumerate(doc.tables):
            did_col = None
            pn_col = None
            header_row_idx = None

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip().upper()
                    if "DID" in text and "对应" in text:
                        did_col = col_idx
                        header_row_idx = row_idx
                    elif text == "DID":
                        did_col = col_idx
                        header_row_idx = row_idx
                    if "件号" in text or "PART" in text.upper():
                        pn_col = col_idx

                if did_col is not None:
                    break

            if did_col is not None:
                for row_idx in range((header_row_idx or 0) + 1, len(table.rows)):
                    row = table.rows[row_idx]
                    did_text = row.cells[did_col].text.strip()

                    m = self.did_pattern.search(did_text)
                    if not m:
                        continue
                    did = m.group(1).upper()

                    if did not in part_numbers:
                        continue

                    if pn_col is not None:
                        actual = row.cells[pn_col].text.strip()
                    else:
                        actual = ""
                        for ci in range(len(row.cells)):
                            if ci != did_col:
                                cell_text = row.cells[ci].text.strip()
                                for prefix in self.pn_prefix_to_did:
                                    if cell_text.upper().startswith(prefix):
                                        actual = cell_text
                                        break
                                if actual:
                                    break

                    if not actual:
                        continue

                    expected = part_numbers[did]
                    status = "OK" if actual == expected else "NG"
                    results.append(CheckResult(
                        file_name=file_name, sheet_name=f"表格{table_idx}",
                        location=f"行{row_idx}", did=did,
                        expected=expected, actual=actual, status=status
                    ))

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    matched_prefix = self._match_pn_prefix(text)
                    if not matched_prefix:
                        continue

                    did = self.pn_prefix_to_did.get(matched_prefix)
                    if not did or did not in part_numbers:
                        continue

                    already_found = any(
                        r.did == did and r.file_name == file_name
                        for r in results
                    )
                    if already_found:
                        continue

                    expected = part_numbers[did]
                    status = "OK" if text == expected else "NG"
                    results.append(CheckResult(
                        file_name=file_name, sheet_name=f"表格{table_idx}",
                        location=f"R{row_idx}C{col_idx}", did=did,
                        expected=expected, actual=text, status=status
                    ))

        return results

    def _check_paragraphs(self, doc, file_name: str, part_numbers: dict[str, str]) -> list[CheckResult]:
        results = []
        paragraphs = doc.paragraphs

        i = 0
        while i < len(paragraphs):
            text = paragraphs[i].text.strip()
            m = self.did_title_pattern.match(text)
            if m:
                did = m.group(1).upper()
                version_found = None
                hex_data_found = None
                location_start = i

                for j in range(i + 1, min(i + 10, len(paragraphs))):
                    next_text = paragraphs[j].text.strip()

                    if j > i + 1 and self.did_title_pattern.match(next_text):
                        break

                    vm = self.version_pattern.search(next_text)
                    if vm:
                        version_found = vm.group(1)

                    hm = self.hex_data_pattern.search(next_text)
                    if hm:
                        hex_data_found = hm.group(1).replace(" ", "").upper()

                if did in part_numbers:
                    expected = part_numbers[did]

                    if did == "F17F" and hex_data_found:
                        actual = hex_data_found
                        expected_clean = expected.upper()
                        status = "OK" if actual == expected_clean else "NG"
                        results.append(CheckResult(
                            file_name=file_name, sheet_name="实测过程",
                            location=f"段落{location_start + 1}-{j}", did=did,
                            expected=expected, actual=actual, status=status
                        ))
                    elif version_found:
                        actual = version_found
                        status = "OK" if actual == expected else "NG"
                        results.append(CheckResult(
                            file_name=file_name, sheet_name="实测过程",
                            location=f"段落{location_start + 1}-{j}", did=did,
                            expected=expected, actual=actual, status=status
                        ))
            i += 1

        found_in_paragraphs = {r.did for r in results}
        has_test_section = len(found_in_paragraphs) > 0

        if has_test_section:
            for did, expected in part_numbers.items():
                if did not in found_in_paragraphs:
                    results.append(CheckResult(
                        file_name=file_name, sheet_name="实测过程",
                        location="-", did=did,
                        expected=expected, actual="-", status="MISSING"
                    ))

        return results

    def _match_pn_prefix(self, val: str) -> Optional[str]:
        val_upper = val.upper()
        for prefix in self.pn_prefix_to_did:
            if val_upper.startswith(prefix):
                return prefix
        return None


# ─── 4. ReportGenerator (Excel 输出) ───

class ReportGenerator:
    HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
    NG_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    MISSING_FILL = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    UNKNOWN_FILL = PatternFill(start_color="B4C6E7", end_color="B4C6E7", fill_type="solid")
    OK_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    THIN_BORDER = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    DETAIL_HEADERS = ["序号", "文件名", "Sheet/区域", "位置", "DID", "期望件号", "实际件号", "结果"]
    COL_WIDTHS = [6, 40, 18, 12, 8, 18, 18, 10]

    def generate(self, results_by_controller: dict[str, list[CheckResult]], output_path: str):
        wb = openpyxl.Workbook()
        first_sheet = True

        for ctrl, results in results_by_controller.items():
            if not results:
                continue
            if first_sheet:
                ws = wb.active
                ws.title = ctrl
                first_sheet = False
            else:
                ws = wb.create_sheet(ctrl)
            self._write_detail_sheet(ws, results)

        ws_summary = wb.create_sheet("汇总统计") if not first_sheet else wb.active
        if first_sheet:
            ws_summary.title = "汇总统计"
        self._write_summary_sheet(ws_summary, results_by_controller)

        wb.save(output_path)

    def _write_detail_sheet(self, ws, results: list[CheckResult]):
        for col, header in enumerate(self.DETAIL_HEADERS, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = self.HEADER_FILL
            cell.font = self.HEADER_FONT
            cell.alignment = Alignment(horizontal='center')
            cell.border = self.THIN_BORDER

        for idx, r in enumerate(results, 1):
            row_data = [idx, r.file_name, r.sheet_name, r.location, r.did, r.expected, r.actual, r.status]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(row=idx + 1, column=col, value=val)
                cell.border = self.THIN_BORDER
                cell.alignment = Alignment(horizontal='center' if col in (1, 5, 8) else 'left')

            if r.status == "NG":
                for c in range(1, 9):
                    ws.cell(row=idx + 1, column=c).fill = self.NG_FILL
            elif r.status == "MISSING":
                for c in range(1, 9):
                    ws.cell(row=idx + 1, column=c).fill = self.MISSING_FILL
            elif r.status == "UNKNOWN":
                for c in range(1, 9):
                    ws.cell(row=idx + 1, column=c).fill = self.UNKNOWN_FILL
            elif r.status == "OK":
                ws.cell(row=idx + 1, column=8).fill = self.OK_FILL

        for i, w in enumerate(self.COL_WIDTHS, 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    def _write_summary_sheet(self, ws, results_by_controller: dict[str, list[CheckResult]]):
        row = 1
        summary_headers = ["控制器", "总检查数", "OK", "NG", "MISSING", "UNKNOWN"]
        for col, h in enumerate(summary_headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.fill = self.HEADER_FILL
            cell.font = self.HEADER_FONT
            cell.alignment = Alignment(horizontal='center')
            cell.border = self.THIN_BORDER

        row = 2
        for ctrl, results in results_by_controller.items():
            if not results:
                continue
            ok_count = sum(1 for r in results if r.status == "OK")
            ng_count = sum(1 for r in results if r.status == "NG")
            miss_count = sum(1 for r in results if r.status == "MISSING")
            unk_count = sum(1 for r in results if r.status == "UNKNOWN")
            for col, val in enumerate([ctrl, len(results), ok_count, ng_count, miss_count, unk_count], 1):
                cell = ws.cell(row=row, column=col, value=val)
                cell.border = self.THIN_BORDER
                cell.alignment = Alignment(horizontal='center')
            row += 1

        for i, w in enumerate([14, 12, 8, 8, 10, 10], 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

        row += 1
        all_ng = [(ctrl, r) for ctrl, rs in results_by_controller.items() for r in rs if r.status == "NG"]
        if all_ng:
            ws.cell(row=row, column=1, value="NG 详情").font = Font(bold=True, size=12)
            row += 1
            ng_headers = ["控制器", "文件名", "Sheet/区域", "DID", "期望件号", "实际件号"]
            for col, h in enumerate(ng_headers, 1):
                cell = ws.cell(row=row, column=col, value=h)
                cell.fill = self.NG_FILL
                cell.font = Font(bold=True)
                cell.border = self.THIN_BORDER
            row += 1
            for ctrl, r in all_ng:
                for col, val in enumerate([ctrl, r.file_name, r.sheet_name, r.did, r.expected, r.actual], 1):
                    cell = ws.cell(row=row, column=col, value=val)
                    cell.border = self.THIN_BORDER
                row += 1

        row += 1
        all_miss = [(ctrl, r) for ctrl, rs in results_by_controller.items() for r in rs if r.status == "MISSING"]
        if all_miss:
            ws.cell(row=row, column=1, value="MISSING 详情").font = Font(bold=True, size=12)
            row += 1
            miss_headers = ["控制器", "文件名", "Sheet/区域", "DID", "期望件号"]
            for col, h in enumerate(miss_headers, 1):
                cell = ws.cell(row=row, column=col, value=h)
                cell.fill = self.MISSING_FILL
                cell.font = Font(bold=True)
                cell.border = self.THIN_BORDER
            row += 1
            for ctrl, r in all_miss:
                for col, val in enumerate([ctrl, r.file_name, r.sheet_name, r.did, r.expected], 1):
                    cell = ws.cell(row=row, column=col, value=val)
                    cell.border = self.THIN_BORDER
                row += 1


# ─── 5. JSON 序列化 ───

def serialize_results(results_by_controller: dict[str, list[CheckResult]]) -> dict:
    all_results = []
    output = {}
    for ctrl, results in results_by_controller.items():
        output[ctrl] = [r.to_dict() for r in results]
        all_results.extend(results)

    ok = sum(1 for r in all_results if r.status == "OK")
    ng = sum(1 for r in all_results if r.status == "NG")
    missing = sum(1 for r in all_results if r.status == "MISSING")
    unknown = sum(1 for r in all_results if r.status == "UNKNOWN")

    return {
        "status": "PASS" if not (ng or missing) else "ISSUES_FOUND",
        "summary": {
            "total": len(all_results),
            "ok": ok,
            "ng": ng,
            "missing": missing,
            "unknown": unknown,
        },
        "results_by_controller": output,
    }


# ─── 6. Skill 入口 ───

def run_check(
    files: list[str],
    controllers: list[str] | None = None,
    pn_file: str | None = None,
    output_format: str = "json",
) -> dict:
    """
    核心校验入口。

    Args:
        files: 报告文件路径列表 (.xlsx/.docx)
        controllers: 控制器名列表，None 则自动检测
        pn_file: 件号对照表路径，默认使用同级 ../part_number/part_number.xlsx
        output_format: "json" | "excel" | "both"

    Returns:
        校验结果 dict
    """
    # 默认件号文件路径
    if pn_file is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        pn_file = os.path.join(script_dir, "..", "part_number", "part_number.xlsx")
        pn_file = os.path.normpath(pn_file)

    # 验证输入
    if not os.path.exists(pn_file):
        return {"error": True, "error_type": "FileNotFoundError", "message": f"件号对照表不存在: {pn_file}"}

    valid_files = []
    for f in files:
        if not os.path.exists(f):
            return {"error": True, "error_type": "FileNotFoundError", "message": f"报告文件不存在: {f}"}
        if not f.endswith((".xlsx", ".docx")):
            return {"error": True, "error_type": "ReportParseError", "message": f"不支持的文件类型: {f}"}
        valid_files.append(f)

    if controllers:
        for c in controllers:
            if c not in CONTROLLER_NAMES:
                return {"error": True, "error_type": "InvalidControllerError",
                        "message": f"未知控制器 '{c}'，有效值: {CONTROLLER_NAMES}"}

    # 加载件号
    all_part_numbers, pn_prefix_to_did = load_part_numbers(pn_file)

    # 初始化 checker
    excel_checker = ExcelChecker(pn_prefix_to_did)
    word_checker = WordChecker(pn_prefix_to_did)

    # 处理文件
    results_by_controller: dict[str, list[CheckResult]] = {}
    for filepath in valid_files:
        fname = os.path.basename(filepath)
        if controllers:
            ctrls = controllers
        else:
            ctrls = resolve_controller(filepath)
            if not ctrls:
                ctrls = list(all_part_numbers.keys())

        for ctrl in ctrls:
            pn = all_part_numbers.get(ctrl, {})
            if filepath.endswith(".xlsx"):
                results = excel_checker.check_file(filepath, fname, pn)
            elif filepath.endswith(".docx"):
                results = word_checker.check_file(filepath, fname, pn)
            else:
                continue
            results_by_controller.setdefault(ctrl, []).extend(results)

    # 构建输出
    output = serialize_results(results_by_controller)

    if output_format in ("excel", "both"):
        result_dir = os.path.join(os.path.dirname(pn_file), "..", "check_result")
        result_dir = os.path.normpath(result_dir)
        os.makedirs(result_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_path = os.path.join(result_dir, f"检查报告_{timestamp}.xlsx")
        ReportGenerator().generate(results_by_controller, excel_path)
        output["excel_report_path"] = excel_path
    else:
        output["excel_report_path"] = None

    return output


# ─── 7. CLI 入口 ───

def main():
    parser = argparse.ArgumentParser(description="ECU Report DID Check (Skill)")
    parser.add_argument("--files", nargs="+", required=True, help="报告文件路径列表")
    parser.add_argument("--controllers", nargs="+", default=None, help="控制器名 (FLSPU/RLSPU/RRSPU/FLSMU/FRSMU)")
    parser.add_argument("--pn-file", default=None, help="件号对照表路径")
    parser.add_argument("--output-format", choices=["json", "excel", "both"], default="json", help="输出格式")
    args = parser.parse_args()

    result = run_check(
        files=args.files,
        controllers=args.controllers,
        pn_file=args.pn_file,
        output_format=args.output_format,
    )

    print(json.dumps(result, ensure_ascii=False, indent=2))

    if result.get("error"):
        sys.exit(1)
    if result.get("status") == "ISSUES_FOUND":
        sys.exit(2)


if __name__ == "__main__":
    main()
